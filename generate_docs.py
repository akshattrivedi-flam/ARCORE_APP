from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_solution_doc():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Solving Handheld & Floating Object Annotations for Objectron', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Introduction ---
    doc.add_heading('1. The Challenge', level=1)
    doc.add_paragraph(
        "Current ARCore dataset collection typically relies on 'Plane Detection' (Horizontal/Vertical). "
        "When an object like a soda can is held in a hand or placed on a transparent surface, "
        "ARCore cannot detect a valid plane, preventing the bounding box from 'sticking' to the object. "
        "To build a truly robust model like the Google Objectron team, we must handle these mid-air scenarios."
    )

    # --- Solution 1: Depth API ---
    doc.add_heading('2. Method A: ARCore Depth API (Recommended)', level=2)
    doc.add_paragraph(
        "The most professional way to handle floating objects is the ARCore Depth API. "
        "Unlike planes, the Depth API provides a distance value for EVERY pixel on the screen."
    )
    list_depth = [
        "How it works: Instead of looking for a floor, we perform a hit-test against the 'Depth Map'.",
        "Advantage: If the user points the crosshair at a can in their hand, ARCore returns the exact 3D point in space where the can surface exists.",
        "Precision: This allows the 9-keypoint bounding box to 'snap' directly onto the handheld object."
    ]
    for item in list_depth:
        doc.add_paragraph(item, style='List Bullet')

    # --- Solution 2: Manual Z-Distance Slider ---
    doc.add_heading('3. Method B: Manual Depth/Z-Offset Slider', level=2)
    doc.add_paragraph(
        "If a device doesn't support the high-resolution Depth API, we can use a mathematical 'Projective' approach."
    )
    doc.add_paragraph(
        "The app can provide a slider (e.g., 20cm to 200cm). As the user moves the slider, "
        "the bounding box moves closer or further away along the camera's forward vector. "
        "This is highly efficient for floating objects because the user can manually 'eye-ball' the fit "
        "regardless of surface detection."
    )

    # --- Solution 3: Instant Placement (Shadow Mapping) ---
    doc.add_heading('4. Method C: Instant Placement API', level=2)
    doc.add_paragraph(
        "ARCore's Instant Placement allows the immediate placement of objects without waiting for plane detection. "
        "It uses a heuristic to guess 3D locations based on screen movement. "
        "While less accurate than Depth, it allows for quick setup of floating annotations."
    )

    # --- Solution 4: Augmented Image Anchoring ---
    doc.add_heading('5. Method D: Augmented Image Tracking', level=2)
    doc.add_paragraph(
        "If the cans have distinct labels (like Coca-Cola or ThumsUp), we can use ARCore's Augmented Images."
    )
    doc.add_paragraph(
        "By providing a high-res image of the can label to the ARSession, the app can automatically "
        "detect the can and attach the 3D bounding box to it at a fixed relative pose. "
        "This is the 'Google-level' way of doing persistent tracking for handheld items."
    )

    # --- Comparison Table ---
    doc.add_heading('6. Comparison of Methods', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Method'
    hdr_cells[1].text = 'Complexity'
    hdr_cells[2].text = 'Accuracy'

    data = [
        ('Depth API', 'Medium', 'High'),
        ('Z-Slider', 'Low', 'Medium (Human-led)'),
        ('Augmented Images', 'High', 'Extremely High'),
        ('Instant Placement', 'Low', 'Low/Medium')
    ]

    for meth, comp, acc in data:
        row_cells = table.add_row().cells
        row_cells[0].text = meth
        row_cells[1].text = comp
        row_cells[2].text = acc

    # --- Suggested Strategy ---
    doc.add_heading('7. Final Recommendation', level=1)
    doc.add_paragraph(
        "For your 45-video dataset, I suggest a Hybrid Approach: ", style='Body Text'
    )
    doc.add_paragraph("1. Implement the Depth API for high-end phones to get automatic surface-snapping.")
    doc.add_paragraph("2. Add a 'Z-Distance' Row in your UI control panel to allow moving the box 'Up/Down/Forward' manually.")
    doc.add_paragraph("3. This combination ensures that whether the can is on a shelf, in a hand, or on the floor, the annotation fits perfectly.")

    # Save
    path = '/home/user/Desktop/Handheld_Object_Annotation_Solutions.docx'
    doc.save(path)
    print(f"Document saved to: {path}")

if __name__ == "__main__":
    create_solution_doc()
